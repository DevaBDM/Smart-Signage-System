"""
Final Project Document Build System
-----------------------------------
This script automates the assembly of Markdown chapters into a professional,
submission-ready Word document. 

Features:
- Configuration-based assembly order (build_config.json)
- Automatic Table of Contents and Page Number refreshing
- Specialized Roman/Arabic section management
- Custom Word style injection via Markdown tags
- Standardized academic table captioning
"""

import os
import subprocess
import re
import json
from docx import Document
from docx.oxml import OxmlElement, ns
from docx.enum.text import WD_ALIGN_PARAGRAPH

# --- Constants & Configuration ---
CONFIG_FILE = "build_config.json"
TEMP_MD_FILE = "FULL_REPORT_TEMP.md"

def load_build_configuration():
    """
    Loads assembly order and filenames from JSON. 
    Creates a default config if one does not exist.
    """
    if not os.path.exists(CONFIG_FILE):
        default_config = {
            "reference_docx": "Final_Document.docx",
            "profiles": [
                {
                    "name": "Main Document",
                    "order": [
                        "TOP_part.md", "CH1.md", "CH2.md", "CH3.md", 
                        "CH4.md", "CH5.md", "CH6.md", "CH7.md", 
                        "Reference.md", "Appendices.md"
                    ],
                    "output_docx": "Final_Document_Built.docx"
                },
                {
                    "name": "Minimal Summary",
                    "order": [
                        "TOP_part_min.md", "CH1_min.md", "CH2_min.md", "CH3_min.md", 
                        "CH4_min.md", "CH5_min.md", "CH6_min.md", "CH7_min.md", 
                        "Reference_min.md", "Appendices_min.md"
                    ],
                    "output_docx": "Final_Document_Summary.docx"
                }
            ]
        }
        with open(CONFIG_FILE, "w", encoding="utf-8") as f:
            json.dump(default_config, f, indent=4)
        return default_config
    
    with open(CONFIG_FILE, "r", encoding="utf-8") as f:
        return json.load(f)

# --- Word Field Helpers ---

def create_word_field(parent_paragraph, instruction):
    """
    Injects a raw Word Field code (e.g., TOC, SEQ, PAGE) into a paragraph.
    """
    run = parent_paragraph.add_run()
    r_element = run._r
    
    # Start the field
    fldChar_begin = OxmlElement('w:fldChar')
    fldChar_begin.set(ns.qn('w:fldCharType'), 'begin')
    r_element.append(fldChar_begin)

    # Insert the instruction
    instrText = OxmlElement('w:instrText')
    instrText.set(ns.qn('xml:space'), 'preserve')
    instrText.text = instruction
    r_element.append(instrText)

    # End the field
    fldChar_end = OxmlElement('w:fldChar')
    fldChar_end.set(ns.qn('w:fldCharType'), 'end')
    r_element.append(fldChar_end)

def add_automated_caption(paragraph, label, section_label=None, reset=False):
    """
    Creates a numbered caption (e.g., Table 4-1: ) using Word Sequence fields.
    If section_label is provided (e.g., 'A'), uses 'Figure A.1:' format.
    Uses 'reset=True' to restart the sequence at 1 (e.g. for the first figure of an appendix).
    """
    paragraph.text = ""
    paragraph.add_run(f"{label} ")
    
    if section_label:
        # Appendix Mode: Use literal letter and reset sequence if requested
        paragraph.add_run(f"{section_label}.")
        
        # We MUST use the standard 'Figure' / 'Table' name so it shows in the LOF/LOT
        if reset:
            instruction = f' SEQ {label} \\* ARABIC \\r 1 '
        else:
            instruction = f' SEQ {label} \\* ARABIC '
        
        create_word_field(paragraph, instruction)
    else:
        # Chapter Mode: Use StyleRef to get the current chapter number from Heading 1
        create_word_field(paragraph, ' STYLEREF 1 \\s ')
        paragraph.add_run("-")
        # Use Seq to get the incrementing number within the chapter
        # \s 1 resets the sequence at each Heading 1
        create_word_field(paragraph, f' SEQ {label} \\* ARABIC \\s 1 ')
    
    paragraph.add_run(": ")

def insert_list_placeholder_field(paragraph, list_type):
    """
    Inserts the active Word field for Table of Contents, Figures, or Tables.
    FIX #2: Added guard for unknown list_type to prevent UnboundLocalError.
    """
    paragraph.text = ""
    if list_type == "TOC":
        # Maps levels 1-3 AND custom styles 'Title' and 'Cover' to the TOC
        instruction = ' TOC \\o "1-3" \\h \\z \\u \\t "Title,1,title,1,Cover,1" '
    elif list_type == "LOF":
        instruction = ' TOC \\h \\z \\c "Figure" '
    elif list_type == "LOT":
        instruction = ' TOC \\h \\z \\c "Table" '
    else:
        print(f"    ! Warning: Unknown list type '{list_type}', skipping field insertion.")
        return
    
    create_word_field(paragraph, instruction)

def add_centered_page_number(paragraph):
    """
    Adds a dynamic PAGE field to a paragraph and centers it.
    """
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    create_word_field(paragraph, "PAGE")

# --- Section & Style Logic ---

def set_section_page_numbering(section, start_value, format_code):
    """
    Configures the numbering format and starting page for a document section.
    """
    sectPr = section._sectPr
    pgNumType = sectPr.xpath('./w:pgNumType')
    
    if not pgNumType:
        pgNumType = OxmlElement('w:pgNumType')
        sectPr.append(pgNumType)
    else:
        pgNumType = pgNumType[0]
    
    if start_value is not None:
        pgNumType.set(ns.qn('w:start'), str(start_value))
    
    if format_code:
        pgNumType.set(ns.qn('w:fmt'), format_code)

def apply_named_style(document, paragraph, style_name, is_header=False):
    """
    Applies a Word style by name (case-insensitive). 
    If style is missing, falls back to default.
    If is_header is True, ensures the paragraph appears in the TOC.
    FIX #5: Log preview handles empty/field paragraphs gracefully.
    """
    target_style = None
    for style in document.styles:
        if style.name.lower() == style_name.lower():
            target_style = style
            break
            
    if target_style:
        try:
            if target_style.type == 1: # Paragraph Style
                paragraph.style = target_style
            elif target_style.type == 2: # Character Style
                if not paragraph.runs:
                    paragraph.add_run(paragraph.text)
                    paragraph.text = ""
                for run in paragraph.runs:
                    run.style = target_style
            # FIX #5: Show a meaningful label even for empty/field-only paragraphs
            preview = paragraph.text[:30] if paragraph.text.strip() else "<field/empty>"
            print(f"    - Applied style '{target_style.name}' to '{preview}'")
        except Exception as e:
            print(f"    ! Error applying style '{style_name}': {e}")
    
    if is_header:
        # Force Outline Level 1 (0 in OpenXML) so it shows in TOC regardless of style
        try:
            pPr = paragraph._element.get_or_add_pPr()
            outlineLvl = OxmlElement('w:outlineLvl')
            outlineLvl.set(ns.qn('w:val'), '0')
            pPr.append(outlineLvl)
        except Exception as e:
            print(f"    ! Warning: Could not set outline level for '{paragraph.text[:30]}': {e}")

# --- Automations ---

def refresh_all_fields_via_word(file_path):
    """
    Uses a background PowerShell session to automate Word field refreshing.
    This replaces the manual Ctrl+A, F9 steps.
    """
    absolute_path = os.path.abspath(file_path)
    print("  + Auto-refreshing TOC and Page Numbers via Word automation...")
    
    ps_script = f"""
    $word = New-Object -ComObject Word.Application
    $word.Visible = $false
    try {{
        $doc = $word.Documents.Open('{absolute_path}')
        $doc.Fields.Update()
        $doc.TablesOfContents | ForEach-Object {{ $_.Update() }}
        $doc.TablesOfFigures | ForEach-Object {{ $_.Update() }}
        $doc.Save()
        $doc.Close()
        Write-Host "SUCCESS"
    }} catch {{
        Write-Host "ERROR: $($_.Exception.Message)"
    }} finally {{
        $word.Quit()
    }}
    """
    
    try:
        process = subprocess.run(["powershell", "-Command", ps_script], capture_output=True, text=True)
        if "SUCCESS" in process.stdout:
            print("    - Refresh complete.")
        else:
            print(f"    ! Auto-refresh failed: {process.stdout.strip()}")
            print("    ! You may need to press F9 manually in Word.")
    except Exception as e:
        print(f"    ! PowerShell automation failed: {e}")

# --- Build Phases ---

def post_process_docx(file_path):
    """
    Final polishing of the Word document: handling placeholders and section rules.
    """
    print(f"Polishing {file_path}...")
    doc = Document(file_path)
    
    # Track Appendix Letter for custom numbering (e.g., Figure A.1)
    current_appendix_label = None
    needs_reset_fig = False
    needs_reset_tbl = False

    # 1. Replace markers with active Word features
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        
        if "CUSTOM_HDR_STYLE_" in text:
            match = re.search(r'CUSTOM_HDR_STYLE_(.*?)_(.*)', text)
            if match:
                style_name, content = match.groups()
                # Update Appendix state
                if "Appendix" in content:
                    app_match = re.search(r'Appendix\s+([A-Z])', content)
                    if app_match:
                        new_label = app_match.group(1)
                        if new_label != current_appendix_label:
                            current_appendix_label = new_label
                            needs_reset_fig = True
                            needs_reset_tbl = True
                else:
                    # If it's a chapter or other major heading, exit Appendix mode
                    if content.startswith("Chapter"):
                        current_appendix_label = None
                
                paragraph.text = content.strip()
                
                # Only force into TOC if it's a major title or cover element
                # 'subtitle' and others should follow the reference doc's style rules
                show_in_toc = style_name.lower() in ["title", "cover"]
                apply_named_style(doc, paragraph, style_name, is_header=show_in_toc)
                continue 

        if text.startswith("AUTO_FIG:"):
            caption_content = text.replace("AUTO_FIG:", "").strip()
            if caption_content:
                add_automated_caption(paragraph, "Figure", current_appendix_label, needs_reset_fig)
                paragraph.add_run(caption_content)
                apply_named_style(doc, paragraph, "Caption")
                needs_reset_fig = False # Reset only once per appendix letter
            else:
                paragraph.text = "" 
                
        elif text.startswith("AUTO_TBL:"):
            caption_content = text.replace("AUTO_TBL:", "").strip()
            if caption_content:
                add_automated_caption(paragraph, "Table", current_appendix_label, needs_reset_tbl)
                paragraph.add_run(caption_content)
                apply_named_style(doc, paragraph, "Caption")
                needs_reset_tbl = False # Reset only once per appendix letter
            else:
                paragraph.text = "" 
                
        elif "CUSTOM_TXT_STYLE_" in text:
            match = re.search(r'CUSTOM_TXT_STYLE_(.*?)_(.*)', text)
            if match:
                style_name, content = match.groups()
                # Remove the style marker while preserving non-text runs (like images)
                for run in paragraph.runs:
                    if "CUSTOM_TXT_STYLE_" in run.text:
                        # Only replace the marker part in the specific run that contains it
                        run.text = run.text.replace(f"CUSTOM_TXT_STYLE_{style_name}_", "")
                        break 
                apply_named_style(doc, paragraph, style_name, is_header=False)
                
        elif "TOCPLACEHOLDER" in text:
            insert_list_placeholder_field(paragraph, "TOC")
        elif "LOFPLACEHOLDER" in text:
            insert_list_placeholder_field(paragraph, "LOF")
        elif "LOTPLACEHOLDER" in text:
            insert_list_placeholder_field(paragraph, "LOT")
                
        # 2. Handle Tab Markers
        if "[TAB]" in text.upper():
            for run in paragraph.runs:
                # Case-insensitive replacement
                run.text = re.sub(r'\[TAB\]', '\t', run.text, flags=re.IGNORECASE)

    # 3. Format Tables
    print(f"  Formatting {len(doc.tables)} tables (Style: Plain Table 2, AutoFit: Contents)...")
    for table in doc.tables:
        # Apply Table Style
        try:
            table.style = 'Plain Table 2'
        except Exception as e:
            print(f"    ! Warning: Could not apply 'Plain Table 2' to a table: {e}")
        
        # AutoFit to Contents
        try:
            tbl = table._element
            tbl_pr = tbl.tblPr
            if tbl_pr is None:
                tbl_pr = OxmlElement('w:tblPr')
                tbl.insert(0, tbl_pr)
            
            # 1. Set Table Width to Auto
            for existing_w in tbl_pr.xpath('./w:tblW'):
                tbl_pr.remove(existing_w)
            tbl_w = OxmlElement('w:tblW')
            tbl_w.set(ns.qn('w:w'), '0')
            tbl_w.set(ns.qn('w:type'), 'auto')
            tbl_pr.append(tbl_w)

            # 2. Set Table Layout to Auto (not fixed)
            for existing_layout in tbl_pr.xpath('./w:tblLayout'):
                tbl_pr.remove(existing_layout)
            tbl_layout = OxmlElement('w:tblLayout')
            tbl_layout.set(ns.qn('w:type'), 'auto')
            tbl_pr.append(tbl_layout)

            # 3. Set each cell width to auto explicitly
            # FIX: Must set tcW w=0 type=auto — not just remove the element.
            # Removing without replacing leaves Word with no signal and causes
            # proportional (equal-column) fallback instead of content-fit.
            for row in table.rows:
                for cell in row.cells:
                    tc_pr = cell._tc.get_or_add_tcPr()
                    for tc_w in tc_pr.xpath('./w:tcW'):
                        tc_pr.remove(tc_w)
                    tc_w_new = OxmlElement('w:tcW')
                    tc_w_new.set(ns.qn('w:w'), '0')
                    tc_w_new.set(ns.qn('w:type'), 'auto')
                    tc_pr.append(tc_w_new)

        except Exception as e:
            print(f"    ! Warning: Could not set AutoFit for a table: {e}")

    # 4. Configure Section Numbering
    print(f"  Found {len(doc.sections)} sections. Applying numbering rules...")
    for index, section in enumerate(doc.sections):
        # Allow independent footers
        section.footer.is_linked_to_previous = False
        
        if index == 0:
            # Cover Page Section
            print("    - Section 0 (Cover): Removing numbering")
            for footer_p in section.footer.paragraphs:
                footer_p.text = ""
            continue
            
        if index == 1:
            # Front Matter (Roman Numerals i, ii, iii)
            print("    - Section 1 (Front Matter): Lower Roman, start at 1")
            set_section_page_numbering(section, 1, 'lowerRoman')
        elif index >= 2:
            # Main Body (Arabic Numerals 1, 2, 3)
            if index == 2:
                print("    - Section 2 (Body Start): Arabic, restart at 1")
                set_section_page_numbering(section, 1, 'decimal')
            else:
                # Subsequent sections continue numbering
                set_section_page_numbering(section, None, 'decimal')
        
        # Add the centered page number field to the footer
        footer_para = section.footer.paragraphs[0] if section.footer.paragraphs else section.footer.add_paragraph()
        footer_para.text = "" 
        add_centered_page_number(footer_para)

    doc.save(file_path)
    
    # Enable fallback for non-automation environments
    doc = Document(file_path)
    settings = doc.settings._element
    update_on_open = settings.find(ns.qn('w:updateFields'))
    if update_on_open is None:
        update_on_open = OxmlElement('w:updateFields')
        update_on_open.set(ns.qn('w:val'), 'true')
        settings.append(update_on_open)
        
    doc.save(file_path)

def preprocess_markdown(text, is_reference_file=False):
    """
    Converts custom signals into intermediate placeholders for post-processing.
    """
    # 0. Handle Citation Linking
    if is_reference_file:
        # Convert [1] or \[1\] at start of line to an anchor: [[1]]{#ref-1}
        # Only matches if it's the main entry (e.g., [1] Author...)
        text = re.sub(r'^\\?\[(\d+)\\?\](?=\s)', r'[[\1]]{#ref-\1}', text, flags=re.MULTILINE)
    else:
        # Convert [1] or \[1\] in text to a link: [[1]](#ref-1)
        text = re.sub(r'\\?\[(\d+)\\?\]', r'[[\1]](#ref-\1)', text)

    # 1. Handle Custom Styles: # Header {.Style} and Text {.Style}
    # Matches any level of header (one or more #)
    text = re.sub(r'^(#+)\s+(.*?)\s+\{\.(.*?)\}', r'\1 CUSTOM_HDR_STYLE_\3_\2', text, flags=re.MULTILINE)
    # FIX #6: Require at least one non-space character (\S) before the style marker
    # to prevent matching blank lines with trailing whitespace.
    text = re.sub(r'^(?!#)(\S.*?)\s+\{\.(.*?)\}', r'CUSTOM_TXT_STYLE_\2_\1', text, flags=re.MULTILINE)

    # 2. Handle Tab Markers (Convert lines with [tab] to raw OpenXML paragraphs)
    processed_lines = []
    for line in text.split('\n'):
        if '[tab]' in line.lower():
            # Extract style if present: "Text [tab] Text {.Style}"
            style_match = re.search(r'\{\.(.*?)\}', line)
            style_xml = ""
            clean_line = line
            if style_match:
                style_name = style_match.group(1)
                style_xml = f'<w:pPr><w:pStyle w:val="{style_name}"/></w:pPr>'
                clean_line = re.sub(r'\s*\{\..*?\}', '', line)
            
            # Split by [tab] and build XML runs
            parts = re.split(r'\[tab\]', clean_line, flags=re.IGNORECASE)
            xml_content = f'```{{=openxml}}\n<w:p>{style_xml}'
            for i, part in enumerate(parts):
                # Clean up bold/italic markers if they are being passed through raw
                part = part.replace('**', '').replace('*', '')
                xml_content += f'<w:r><w:t xml:space="preserve">{part}</w:t></w:r>'
                if i < len(parts) - 1:
                    xml_content += '<w:r><w:tab/></w:r>'
            xml_content += '</w:p>\n```'
            processed_lines.append(xml_content)
        else:
            processed_lines.append(line)
    text = '\n'.join(processed_lines)

    # 3. Handle Table Captions (Convert signals and ensure blank line for Pandoc)
    lines = text.split('\n')
    processed_lines = []
    for line in lines:
        stripped = line.strip()
        if stripped.lower().startswith('<caption>') or stripped.startswith('<!-- CAPTION:'):
            caption_text = ""
            if stripped.startswith('<!--'):
                match = re.search(r'<!--\s*CAPTION:\s*(.*?)\s*-->', stripped, re.IGNORECASE)
                caption_text = match.group(1) if match else ""
            else:
                caption_text = re.sub(r'</?caption>', '', stripped, flags=re.IGNORECASE).strip()
            
            clean_name = re.sub(r'^(Table)\s*[A-Z0-9\.\-]*[:\s]*', '', caption_text, flags=re.IGNORECASE).strip()
            if clean_name:
                processed_lines.append(f'AUTO_TBL: {clean_name}')
                processed_lines.append('') # Crucial spacing
        else:
            processed_lines.append(line)
    text = '\n'.join(processed_lines)

    # 4. Handle Figure/Image Blocks
    figure_pattern = re.compile(r'<figure>.*?(<img[^>]+>).*?<figcaption>(.*?)</figcaption>.*?</figure>', re.DOTALL | re.IGNORECASE)
    def figure_match_callback(match):
        img_tag = match.group(1)
        caption_html = match.group(2)
        caption_text = re.sub(r'<[^>]+>', '', caption_html).strip()
        clean_caption = re.sub(r'^(Figure|Fig)\s*[A-Z0-9\.\-]*[:\s]*', '', caption_text, flags=re.IGNORECASE).strip()
        # FIX #3: Guard against missing src attribute to prevent AttributeError crash
        src_match = re.search(r'src="([^"]+)"', img_tag)
        if not src_match:
            return match.group(0)  # Return original unchanged if no src found
        src = src_match.group(1)
        width = re.search(r'width:([\d\.]+)in', img_tag)
        height = re.search(r'height:([\d\.]+)in', img_tag)
        w_arg = f'width={width.group(1)}in' if width else ""
        h_arg = f'height={height.group(1)}in' if height else ""
        return f'\n![AUTO_FIG: {clean_caption}]({src}){{{w_arg} {h_arg}}}\n'
    text = figure_pattern.sub(figure_match_callback, text)

    standalone_img_pattern = re.compile(r'<img\s+[^>]*src="([^"]+)"[^>]*>', re.IGNORECASE | re.DOTALL)
    def img_match_callback(match):
        full_tag = match.group(0)
        src = match.group(1)
        alt_match = re.search(r'alt="([^"]*)"', full_tag)
        alt_text = alt_match.group(1) if alt_match else ""
        if "width" in full_tag or "height" in full_tag:
            clean_alt = re.sub(r'^(Figure|Fig)\s*[A-Z0-9\.\-]*[:\s]*', '', alt_text, flags=re.IGNORECASE).strip()
            if clean_alt:
                alt_text = "AUTO_FIG: " + clean_alt
        width = re.search(r'width:([\d\.]+)in', full_tag)
        height = re.search(r'height:([\d\.]+)in', full_tag)
        w_arg = f'width={width.group(1)}in' if width else ""
        h_arg = f'height={height.group(1)}in' if height else ""
        return f'![{alt_text}]({src}){{{w_arg} {h_arg}}}'
    text = standalone_img_pattern.sub(img_match_callback, text)

    # 5. Handle Page and Section Breaks
    text = text.replace('\\newpage', '\n```{=openxml}\n<w:p><w:r><w:br w:type="page"/></w:r></w:p>\n```\n')
    text = text.replace('<!-- SECTION_BREAK_ROMAN -->', '\n```{=openxml}\n<w:p><w:pPr><w:sectPr><w:pgNumType w:fmt="romanLower" w:start="1"/><w:type w:val="nextPage"/></w:sectPr></w:pPr></w:p>\n```\n')
    text = text.replace('<!-- SECTION_BREAK_ARABIC -->', '\n```{=openxml}\n<w:p><w:pPr><w:sectPr><w:pgNumType w:fmt="decimal" w:start="1"/><w:type w:val="nextPage"/></w:sectPr></w:pPr></w:p>\n```\n')
    text = text.replace('<!-- SECTION_BREAK -->', '\n```{=openxml}\n<w:p><w:pPr><w:sectPr><w:type w:val="nextPage" /></w:sectPr></w:pPr></w:p>\n```\n')

    # 6. Insert Automated List Placeholders
    text = re.sub(r'# (?:CUSTOM_HDR_STYLE_.*?_)?Table of Contents', r'\g<0>\n\nTOCPLACEHOLDER\n', text, flags=re.IGNORECASE)
    text = re.sub(r'# (?:CUSTOM_HDR_STYLE_.*?_)?List of Figures', r'\g<0>\n\nLOFPLACEHOLDER\n', text, flags=re.IGNORECASE)
    text = re.sub(r'# (?:CUSTOM_HDR_STYLE_.*?_)?List of Tables', r'\g<0>\n\nLOTPLACEHOLDER\n', text, flags=re.IGNORECASE)
    
    return text

def execute_build():
    """
    Main orchestration function for the build system.
    """
    config = load_build_configuration()
    profiles = config.get("profiles", [])
    style_reference = config.get("reference_docx", "Final_Document.docx")
    
    if not profiles:
        print("  ! Error: No build profiles found in configuration.")
        return

    print(f"--- Starting Document Build System ({len(profiles)} profiles) ---")
    
    for profile in profiles:
        name = profile.get("name", "Unnamed")
        assembly_order = profile.get("order", [])
        output_docx = profile.get("output_docx", f"{name}.docx")
        
        print(f"\n>> Building Profile: {name}")
        print(f"   Target: {output_docx}")
        
        try:
            # Step 1: Combine and Preprocess
            with open(TEMP_MD_FILE, "w", encoding="utf-8") as output_markdown:
                for filename in assembly_order:
                    if os.path.exists(filename):
                        print(f"     + Processing {filename}")
                        with open(filename, "r", encoding="utf-8") as source_file:
                            source_text = source_file.read()
                            # Check if this is a reference file to apply anchor logic
                            is_ref = "reference" in filename.lower()
                            processed_text = preprocess_markdown(source_text, is_reference_file=is_ref)
                            output_markdown.write(processed_text + "\n\n")
                    else:
                        print(f"     ! Warning: {filename} missing, skipping.")
            
            # Step 2: Convert to Word via Pandoc
            print("     Converting to Word via Pandoc...")
            pandoc_cmd = [
                "pandoc", TEMP_MD_FILE, 
                "-o", output_docx, 
                "--from", "markdown+implicit_figures+table_captions+raw_attribute", 
                "--resource-path", ".", 
                "--standalone"
            ]
            
            if os.path.exists(style_reference):
                pandoc_cmd.extend(["--reference-doc", style_reference])
            
            # FIX #7: Capture stderr so Pandoc errors are visible in the log
            conversion_result = subprocess.run(pandoc_cmd, capture_output=True, text=True)
            
            if conversion_result.returncode == 0:
                # Step 3: Polish in Python (Styles, Sections, Footers)
                post_process_docx(output_docx)
                
                # Step 4: Final Automation (Background Word Field Update)
                refresh_all_fields_via_word(output_docx)
                
                print(f"   - SUCCESS: Created {output_docx}")
            else:
                print(f"   ! ERROR: Pandoc conversion failed for {name}.")
                # FIX #7: Print Pandoc's actual error message
                if conversion_result.stderr:
                    print(f"   ! Pandoc said: {conversion_result.stderr.strip()}")
                
        except Exception as e:
            print(f"   ! CRITICAL ERROR building {name}: {e}")
        finally:
            # FIX #4: Always clean up temp file, even if post_process_docx or
            # refresh_all_fields_via_word raises an exception mid-build.
            if os.path.exists(TEMP_MD_FILE):
                os.remove(TEMP_MD_FILE)

    print("\n--- Build System Finished ---")

if __name__ == "__main__":
    execute_build()
